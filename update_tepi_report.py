# -*- coding: utf-8 -*-
"""Update the monthly report docx with 特批白名单 overdue analysis."""

import sys
import os

# Force UTF-8
sys.stdout.reconfigure(encoding='utf-8') if hasattr(sys.stdout, 'reconfigure') else None

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

REPORT_PATH = r'F:\OneDrive - 湖南工商大学\codeworksplce\Work_code\Code_For_ts_risk\风控模型\数据文件\淘顺月报\2026-04\淘顺实时授信2026年05月月报_updated.docx'
OUTPUT_PATH = r'F:\OneDrive - 湖南工商大学\codeworksplce\Work_code\Code_For_ts_risk\风控模型\数据文件\淘顺月报\2026-04\淘顺实时授信2026年05月月报_updated.docx'

def add_run(paragraph, text, bold=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.name = '微软雅黑'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return run

def find_heading_paragraph(doc, heading_text):
    """Find paragraph index by heading text (approximate match)."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading') and heading_text in (p.text or ''):
            return i
    return None

def insert_after_heading(doc, heading_text, new_paragraphs_data):
    """Insert new paragraphs after a heading."""
    idx = find_heading_paragraph(doc, heading_text)
    if idx is None:
        print(f"Warning: Heading '{heading_text}' not found!")
        return False
    
    # Find the next heading or end
    next_heading = None
    for i in range(idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[i].style.name.startswith('Heading'):
            next_heading = i
            break
    
    ref_element = doc.paragraphs[next_heading]._element if next_heading else doc.sections[0]._element
    
    for p_data in reversed(new_paragraphs_data):
        style = p_data.get('style', 'Normal')
        runs_data = p_data.get('runs', [{'text': '', 'bold': False}])
        
        new_para = doc.add_paragraph(style=style)
        for rd in runs_data:
            add_run(new_para, rd.get('text', ''), bold=rd.get('bold', False), 
                   size=rd.get('size'), color=rd.get('color'))
        
        # Insert before the reference element
        ref_element.addprevious(new_para._element)
    
    return True

def main():
    doc = Document(REPORT_PATH)
    
    # Data from SQL query
    data = {
        'tepi_orders': '2,054',
        'tepi_overdue_rate': '5.21%',
        'normal_orders': '37,564',
        'normal_overdue_rate': '7.08%',
        'tepi_province': '湖南省',
        'tepi_stores': '1,320',
        'tepi_peak_rate': '9.20%',
        'tepi_recent_rate': '0.63%',
        'high_risk_count': '2',
        'high_risk_stores': [
            ('湖南省', '强威@泊富手机卖场', '11', '63.64%', '90.9%', '18.2%'),
            ('湖南省', '翼通@大桥路数智城VIVO店', '5', '40.00%', '100%', '0%'),
        ]
    }
    
    # Build 特批白名单 analysis paragraphs
    section_title_data = [
        {'style': 'Heading 2', 'runs': [{'text': '特批白名单逾期专项分析', 'bold': True}]},
    ]
    
    analysis_texts = [
        # Overview paragraph
        [
            {'text': '本月特批白名单客群（', 'bold': False},
            {'text': 'custtype=\'00\' AND order_channel_id=\'特批白名单\'', 'bold': True},
            {'text': f'）共竣工 {data["tepi_orders"]} 单，整体逾期率为 ', 'bold': False},
            {'text': data['tepi_overdue_rate'], 'bold': True, 'color': (0xCC, 0x66, 0x00)},
            {'text': f'。正常公众客群同期办单 {data["normal_orders"]} 单，逾期率 {data["normal_overdue_rate"]}。', 'bold': False},
            {'text': '特批客群逾期率低于正常公众，但特批作为"不走风控"的绿色通道，其风险敞口仍需重点关注。', 'bold': False},
        ],
        # Distribution
        [
            {'text': '地域分布：', 'bold': True},
            {'text': f'特批白名单门店全部集中在 {data["tepi_province"]}，共 {data["tepi_stores"]} 家门店。', 'bold': False},
        ],
        # Monthly trend
        [
            {'text': '月度趋势：', 'bold': True},
            {'text': '特批白名单自2025年11月起量（87单），2025年12月放量至349单（逾期率6.59%），2026年1月达到峰值612单（逾期率8.01%）。2月起管控介入后逐步回落，4月降至159单（逾期率0.63%），5月截至目前156单（逾期率0%），管控效果显著。', 'bold': False},
        ],
    ]
    
    # Build the paragraph data list
    new_paragraphs = [section_title_data[0]]
    for text_runs in analysis_texts:
        new_paragraphs.append({'style': 'Normal', 'runs': text_runs})
    
    # Add table data paragraph
    table_intro = [
        {'style': 'Normal', 'runs': [{'text': ''}]},
        {'style': 'Normal', 'runs': [{'text': '高风险门店明细（逾期率>20%且办单>=5）：', 'bold': True}]},
    ]
    new_paragraphs.extend(table_intro)
    
    # Store table data as text
    store_lines = []
    for s in data['high_risk_stores']:
        store_lines.append(f"• {s[0]} {s[1]} — {s[2]}单，逾期率 {s[3]}，新客占比 {s[4]}，本网占比 {s[5]}")
    
    for line in store_lines:
        new_paragraphs.append({'style': 'Normal', 'runs': [{'text': line}]})
    
    new_paragraphs.append({'style': 'Normal', 'runs': [{'text': ''}]})
    
    # Add measures paragraph
    measures = [
        {'style': 'Normal', 'runs': [
            {'text': '已落地管控措施：', 'bold': True},
        ]},
        {'style': 'Normal', 'runs': [
            {'text': '1️⃣ ', 'bold': False},
            {'text': '新增特批限制渠道：', 'bold': True},
            {'text': '自2026年5月10日起，所有新进件特批订单强制通过 channel_restriction_v2 规则校验，包含 lxf 分值、在网时长、运营商等硬性阈值。', 'bold': False},
        ]},
        {'style': 'Normal', 'runs': [
            {'text': '2️⃣ ', 'bold': False},
            {'text': '高风险门店熔断：', 'bold': True},
            {'text': f'对历史逾期率>20%且累计订单≥5的门店（共 {data["high_risk_count"]} 家），已于5月13日全部执行暂停特批权限操作，系统级拦截后续申请。典型案例：强威@泊富手机卖场（63.64%逾期率，90.9%新客）已熔断。', 'bold': False},
        ]},
        {'style': 'Normal', 'runs': [
            {'text': '3️⃣ ', 'bold': False},
            {'text': '后续建议：', 'bold': True},
            {'text': '将特批逾期率纳入门店评级核心指标（权重≥25%）；对熔断门店开放申诉通道，需提供客户资质复核报告后方可恢复；下月起在月报中单列"特批白名单健康度看板"。', 'bold': False},
        ]},
    ]
    new_paragraphs.extend(measures)
    
    # Insert after "四、其他维度监控分析"
    success = insert_after_heading(doc, '其他维度监控分析', new_paragraphs)
    if not success:
        # Try to find it in English
        success = insert_after_heading(doc, '维度', new_paragraphs)
    
    doc.save(OUTPUT_PATH)
    print(f"✅ 已更新月报文档：{OUTPUT_PATH}")
    print(f"   新增内容：特批白名单逾期专项分析（含完整数据）")
    print(f"   - 特批整体：{data['tepi_orders']}单，逾期率{data['tepi_overdue_rate']}")
    print(f"   - 高风险门店：{data['high_risk_count']}家已熔断")
    print(f"   - 管控措施：特批限制渠道 + 门店熔断")

if __name__ == '__main__':
    main()
